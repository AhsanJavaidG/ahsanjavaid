"""
Generic AI resume generator for Ahsan Javaid.

This script reads:
- CV from PDF
- Job description from a text/markdown/pdf file (or inline config text)

Then it uses Gemini to tailor a role-specific resume and writes a DOCX file for review.

Config file: resume_ai_config.json
"""

from __future__ import annotations

import json
import os
import re
from pathlib import Path
from typing import Any, Dict, List

import google.generativeai as genai
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from pypdf import PdfReader


SCRIPT_DIR = Path(__file__).resolve().parent
CONFIG_PATH = SCRIPT_DIR / "resume_ai_config.json"
TEAL = RGBColor(0x00, 0x7A, 0x87)


def load_config(path: Path) -> Dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Config file not found: {path}")
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def resolve_api_key(config: Dict[str, Any]) -> str:
    direct_key = str(config.get("gemini_api_key", "")).strip()
    if direct_key:
        return direct_key

    env_name = str(config.get("gemini_api_key_env", "GEMINI_API_KEY")).strip()
    env_value = os.getenv(env_name, "").strip()
    if env_value:
        return env_value

    raise RuntimeError(
        "Gemini API key not found. Set gemini_api_key in resume_ai_config.json "
        f"or set the environment variable {env_name}."
    )


def extract_pdf_text(pdf_path: Path) -> str:
    reader = PdfReader(str(pdf_path))
    chunks: List[str] = []
    for page in reader.pages:
        chunks.append(page.extract_text() or "")

    text = "\n".join(chunks)
    text = text.replace("\u25a1", " ")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def load_job_description(config: Dict[str, Any]) -> str:
    inline_jd = str(config.get("job_description_text", "")).strip()
    if inline_jd:
        return inline_jd

    jd_path_value = str(config.get("job_description_path", "")).strip()
    if not jd_path_value:
        raise ValueError(
            "No job description provided. Set job_description_path or job_description_text in resume_ai_config.json."
        )

    jd_path = SCRIPT_DIR / jd_path_value
    if not jd_path.exists():
        raise FileNotFoundError(f"Job description file not found: {jd_path}")

    if jd_path.suffix.lower() == ".pdf":
        return extract_pdf_text(jd_path)

    return read_text_file(jd_path)


def build_prompt(cv_text: str, job_description: str, config: Dict[str, Any]) -> str:
    target_role = str(config.get("target_role", "Senior Product Designer"))
    target_company = str(config.get("target_company", "")).strip() or "General"
    tone = str(config.get("tone", "concise, executive, impact-oriented"))
    max_bullets = int(config.get("max_bullets_per_role", 4))

    return f"""
You are an expert resume writer and career strategist.

Task:
- Read the candidate CV and the target job description.
- Tailor the resume for best relevance to the job description while staying fully factual.

Target role: {target_role}
Target company: {target_company}
Tone: {tone}
Max bullets per role: {max_bullets}

Return valid JSON only with this schema:
{{
  "full_name": "",
  "headline": "",
  "location": "",
  "mobility": "",
  "citizenship": "",
  "contact": {{
    "email": "",
    "phone_primary": "",
    "phone_secondary": "",
    "linkedin": "",
    "portfolio": ""
  }},
  "summary": "",
  "skills": ["", ""],
  "experience": [
    {{
      "role": "",
      "company": "",
      "location": "",
      "period": "",
      "bullets": ["", ""]
    }}
  ],
  "education": [
    {{
      "degree": "",
      "institution": "",
      "period": ""
    }}
  ],
  "languages": ["", ""]
}}

Hard constraints:
- Do not invent employers, education, or certifications not in CV.
- Keep all claims grounded in CV evidence.
- Reorder and rewrite bullets to align with the target job description.
- Use measurable outcomes where present in CV.

Candidate CV:
{cv_text}

Job Description:
{job_description}
""".strip()


def strip_code_fences(text: str) -> str:
    cleaned = text.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
    cleaned = re.sub(r"\s*```$", "", cleaned)
    return cleaned.strip()


def parse_resume_json(model_text: str) -> Dict[str, Any]:
    cleaned = strip_code_fences(model_text)
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, re.DOTALL)
        if not match:
            raise ValueError("Gemini response does not contain valid JSON.")
        return json.loads(match.group(0))


def generate_resume_json(prompt: str, api_key: str, model_name: str) -> Dict[str, Any]:
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(model_name)
    response = model.generate_content(prompt)

    if not getattr(response, "text", ""):
        raise RuntimeError("Gemini returned an empty response.")

    return parse_resume_json(response.text)


def add_section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = TEAL


def build_docx_resume(resume: Dict[str, Any], out_path: Path) -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.1)
        section.bottom_margin = Cm(1.1)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    name = str(resume.get("full_name", "AHSAN JAVAID"))
    headline = str(resume.get("headline", "Senior Product Designer"))

    title = doc.add_paragraph()
    r = title.add_run(name.upper())
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = TEAL

    doc.add_paragraph(headline)

    contact = resume.get("contact", {}) or {}
    location = str(resume.get("location", "Malmo, Sweden"))
    mobility = str(resume.get("mobility", "SE / DK / Hybrid / Remote"))
    citizenship = str(resume.get("citizenship", "EU (Danish)"))
    contact_line = (
        f"{location} | Mobility: {mobility} | Citizenship: {citizenship} | "
        f"{contact.get('phone_primary', '')} | {contact.get('email', '')}"
    )
    doc.add_paragraph(contact_line)

    add_section_heading(doc, "Professional Summary")
    doc.add_paragraph(str(resume.get("summary", "")))

    add_section_heading(doc, "Core Skills")
    for skill in resume.get("skills", []) or []:
        doc.add_paragraph(str(skill), style="List Bullet")

    add_section_heading(doc, "Professional Experience")
    for item in resume.get("experience", []) or []:
        role = str(item.get("role", ""))
        company = str(item.get("company", ""))
        loc = str(item.get("location", ""))
        period = str(item.get("period", ""))

        header = f"{role} - {company}"
        if loc:
            header += f", {loc}"
        if period:
            header += f" ({period})"

        hp = doc.add_paragraph()
        hr = hp.add_run(header)
        hr.bold = True
        hr.font.color.rgb = TEAL

        for bullet in item.get("bullets", []) or []:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(str(bullet))

    add_section_heading(doc, "Education")
    for edu in resume.get("education", []) or []:
        degree = str(edu.get("degree", ""))
        institution = str(edu.get("institution", ""))
        period = str(edu.get("period", ""))
        doc.add_paragraph(f"{degree} - {institution} ({period})")

    add_section_heading(doc, "Languages")
    for lang in resume.get("languages", []) or []:
        doc.add_paragraph(str(lang), style="List Bullet")

    doc.save(str(out_path))


def main() -> None:
    config = load_config(CONFIG_PATH)

    input_pdf = SCRIPT_DIR / str(config.get("input_pdf", "CV-M-AhsanJavaid.pdf"))
    if not input_pdf.exists():
        raise FileNotFoundError(f"Input CV PDF not found: {input_pdf}")

    api_key = resolve_api_key(config)
    model_name = str(config.get("gemini_model", "gemini-1.5-pro"))

    cv_text = extract_pdf_text(input_pdf)
    job_description = load_job_description(config)
    prompt = build_prompt(cv_text, job_description, config)

    resume_json = generate_resume_json(prompt, api_key, model_name)

    output_docx = SCRIPT_DIR / str(config.get("output_docx", "Ahsan_Javaid_Tailored_Resume.docx"))
    build_docx_resume(resume_json, output_docx)

    print(f"DOCX generated for review: {output_docx}")


if __name__ == "__main__":
    main()
